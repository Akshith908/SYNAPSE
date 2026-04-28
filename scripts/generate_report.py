from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "report_assets"
OUTPUT = ROOT / "SYNAPSE_Course_Based_Project_Report.docx"

PROJECT_TITLE = "SYNAPSE Human Connectome Explorer"
PROJECT_SUBTITLE = "MongoDB-backed Interactive Connectome Simulation Dashboard"
ACADEMIC_YEAR = "2026"
DEPARTMENT = "Department of Information Technology"
COLLEGE = "Vasavi College of Engineering (Autonomous)"
COLLEGE_LINE_2 = "Accredited by NAAC with 'A++' Grade"
COLLEGE_LINE_3 = "(Affiliated to Osmania University and Approved by AICTE)"
COLLEGE_LOCATION = "Ibrahimbagh, Hyderabad - 500031"
TEAM_MEMBERS = [
    ("<Student Name 1>", "<Hall Ticket No. 1>"),
    ("<Student Name 2>", "<Hall Ticket No. 2>"),
    ("<Student Name 3>", "<Hall Ticket No. 3>"),
]
FACULTY_COORDINATORS = ["<Faculty Coordinator 1>", "<Faculty Coordinator 2>"]
GITHUB_LINK = "<Add GitHub repository URL here>"
TODAY = datetime.now().strftime("%d %B %Y")

REFERENCE_ITEMS = [
    ("R1", "MDN Web Docs, Canvas API, https://developer.mozilla.org/en-US/docs/Web/API/Canvas_API, accessed on 23 April 2026."),
    ("R2", "MDN Web Docs, Basic usage of canvas, https://developer.mozilla.org/en-US/docs/Web/API/Canvas_API/Tutorial/Basic_usage, accessed on 23 April 2026."),
    ("R3", "MDN Web Docs, <canvas> element reference, https://developer.mozilla.org/docs/Web/HTML/Reference/Elements/canvas, accessed on 23 April 2026."),
    ("R4", "Express.js, Serving static files in Express, https://expressjs.com/en/starter/static-files, accessed on 23 April 2026."),
    ("R5", "MongoDB Documentation, Node.js Driver, https://www.mongodb.com/developer/technologies/nodejs/, accessed on 23 April 2026."),
    ("R6", "MongoDB Documentation, MongoDB Manual, https://www.mongodb.com/docs/manual/, accessed on 23 April 2026."),
    ("R7", "MongoDB Documentation, Manual Reference, https://www.mongodb.com/docs/manual/reference/, accessed on 23 April 2026."),
    ("R8", "webpack, Getting Started, https://webpack.js.org/guides/getting-started/, accessed on 23 April 2026."),
    ("R9", "webpack, Concepts, https://webpack.js.org/concepts/, accessed on 23 April 2026."),
    ("R10", "npm, dotenv package documentation, https://www.npmjs.com/package/dotenv, accessed on 23 April 2026."),
    ("R11", "Node.js Documentation, File system module, https://nodejs.org/api/fs.html, accessed on 23 April 2026."),
]


def ensure_assets():
    ASSET_DIR.mkdir(exist_ok=True)
    create_architecture_diagram(ASSET_DIR / "architecture_diagram.png")
    create_activity_diagram(ASSET_DIR / "activity_diagram.png")


def get_font(size, bold=False):
    candidates = [
        "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def create_box(draw, box, text, fill, outline, text_color=(20, 24, 32), bold=False):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    font = get_font(28 if bold else 24, bold=bold)
    x0, y0, x1, y1 = box
    text_box = draw.textbbox((0, 0), text, font=font)
    text_w = text_box[2] - text_box[0]
    text_h = text_box[3] - text_box[1]
    draw.text(
        ((x0 + x1 - text_w) / 2, (y0 + y1 - text_h) / 2 - 3),
        text,
        font=font,
        fill=text_color,
    )


def create_arrow(draw, start, end, color):
    draw.line([start, end], fill=color, width=5)
    ax, ay = end
    bx, by = start
    if abs(ax - bx) >= abs(ay - by):
        direction = 1 if ax >= bx else -1
        draw.polygon(
            [(ax, ay), (ax - 18 * direction, ay - 10), (ax - 18 * direction, ay + 10)],
            fill=color,
        )
    else:
        direction = 1 if ay >= by else -1
        draw.polygon(
            [(ax, ay), (ax - 10, ay - 18 * direction), (ax + 10, ay - 18 * direction)],
            fill=color,
        )


def create_architecture_diagram(path):
    image = Image.new("RGB", (1800, 1000), (246, 248, 252))
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    draw.text((60, 40), "SYNAPSE Architecture Diagram", fill=(20, 27, 45), font=title_font)

    boxes = {
        "browser": (80, 180, 470, 360),
        "controller": (540, 180, 930, 360),
        "renderer": (1000, 100, 1390, 280),
        "model": (1000, 330, 1390, 510),
        "api": (1460, 180, 1750, 360),
        "db": (1460, 520, 1750, 700),
    }

    create_box(draw, boxes["browser"], "Single Page Dashboard\nHTML + CSS + JS", (224, 242, 254), (59, 130, 246), bold=True)
    create_box(draw, boxes["controller"], "App Controller\nEvent Handling + State", (236, 253, 245), (16, 185, 129), bold=True)
    create_box(draw, boxes["renderer"], "Canvas Renderer\nGraph + Matrix + Labels", (254, 249, 195), (245, 158, 11))
    create_box(draw, boxes["model"], "Connectome Model\nSeeded Graph + Path Logic", (243, 232, 255), (168, 85, 247))
    create_box(draw, boxes["api"], "Express API\n/bootstrap + /simulations", (254, 226, 226), (239, 68, 68), bold=True)
    create_box(draw, boxes["db"], "MongoDB\nregions / networks /\ndiseaseProfiles / pathways /\nsimulationSnapshots", (229, 231, 235), (75, 85, 99), bold=True)

    create_arrow(draw, (470, 270), (540, 270), (59, 130, 246))
    create_arrow(draw, (930, 230), (1000, 190), (16, 185, 129))
    create_arrow(draw, (930, 310), (1000, 410), (16, 185, 129))
    create_arrow(draw, (1390, 190), (1460, 250), (245, 158, 11))
    create_arrow(draw, (1390, 420), (1460, 290), (168, 85, 247))
    create_arrow(draw, (1605, 360), (1605, 520), (239, 68, 68))

    note_font = get_font(24)
    draw.text((80, 800), "Data flow: dashboard actions -> controller -> graph/model logic -> API -> MongoDB -> dashboard refresh", fill=(55, 65, 81), font=note_font)
    image.save(path)


def create_activity_diagram(path):
    image = Image.new("RGB", (1700, 1200), (250, 251, 255))
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, bold=True)
    draw.text((60, 40), "SYNAPSE Activity Diagram", fill=(20, 27, 45), font=title_font)

    steps = [
        ("Open dashboard", (650, 130, 1050, 240), (224, 242, 254), (59, 130, 246)),
        ("Load sample or MongoDB data", (620, 300, 1080, 410), (236, 253, 245), (16, 185, 129)),
        ("Select region / choose profile", (610, 470, 1090, 580), (254, 249, 195), (245, 158, 11)),
        ("Fire signal / trace path /\nlesion region", (620, 640, 1080, 780), (243, 232, 255), (168, 85, 247)),
        ("Update graph, metrics,\nlogs and labels", (620, 860, 1080, 1000), (254, 226, 226), (239, 68, 68)),
    ]

    for text, box, fill, outline in steps:
        create_box(draw, box, text, fill, outline, bold=True)

    for index in range(len(steps) - 1):
        current = steps[index][1]
        nxt = steps[index + 1][1]
        create_arrow(draw, ((current[0] + current[2]) // 2, current[3]), ((nxt[0] + nxt[2]) // 2, nxt[1]), (75, 85, 99))

    save_box = (120, 640, 500, 780)
    create_box(draw, save_box, "Optional: Save Snapshot", (229, 231, 235), (75, 85, 99))
    create_arrow(draw, (620, 710), (500, 710), (75, 85, 99))
    create_arrow(draw, (310, 780), (310, 950), (75, 85, 99))
    create_arrow(draw, (310, 950), (620, 950), (75, 85, 99))

    font = get_font(24)
    draw.text((1180, 685), "MongoDB stores snapshot summary\nfor later retrieval in the same dashboard.", fill=(55, 65, 81), font=font)
    image.save(path)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = " "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE")


def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)

    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])


def add_centered(document, text, size=12, bold=False, color=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(1.27)
        paragraph.add_run(item)


def add_code_block(document, lines):
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(1.27)
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def add_table(document, headers, rows, col_widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        para = header_cells[index].paragraphs[0]
        para.paragraph_format.first_line_indent = Cm(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            para = row_cells[index].paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.add_run(value)
            row_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for cell, width in zip(row.cells, col_widths):
                cell.width = width

    document.add_paragraph()
    return table


def add_image(document, path, caption, width=Inches(6.2)):
    if path.exists():
        document.add_picture(str(path), width=width)
        paragraph = document.paragraphs[-1]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        run = cap.add_run(caption)
        run.italic = True
    else:
        add_body(document, f"{caption} (Image file missing: {path.name})")


def add_cover_page(document):
    add_centered(document, COLLEGE, 16, True)
    add_centered(document, COLLEGE_LINE_2, 12, True)
    add_centered(document, COLLEGE_LINE_3, 12)
    add_centered(document, COLLEGE_LOCATION, 12)
    document.add_paragraph()
    add_centered(document, PROJECT_TITLE, 20, True, (20, 27, 45))
    add_centered(document, PROJECT_SUBTITLE, 14, True)
    document.add_paragraph()
    add_centered(document, "A", 14, True)
    add_centered(document, "Course Based Project Report", 16, True)
    add_centered(document, "Submitted in partial fulfilment of the requirements for the award of the degree of", 12)
    add_centered(document, "BACHELOR OF ENGINEERING", 14, True)
    add_centered(document, "IN", 12, True)
    add_centered(document, "INFORMATION TECHNOLOGY", 14, True)
    document.add_paragraph()
    add_centered(document, "By", 12, True)
    for name, hall_ticket in TEAM_MEMBERS:
        add_centered(document, f"{name}    {hall_ticket}", 12)
    document.add_paragraph()
    add_centered(document, DEPARTMENT, 14, True)
    add_centered(document, COLLEGE, 14, True)
    add_centered(document, ACADEMIC_YEAR, 12, True)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.first_line_indent = Cm(0)
    run = note.add_run("Replace placeholder names, hall ticket numbers, and faculty details before final submission.")
    run.italic = True
    run.font.size = Pt(11)
    document.add_page_break()


def add_declaration_page(document):
    add_centered(document, COLLEGE, 14, True)
    add_centered(document, COLLEGE_LINE_2, 12, True)
    add_centered(document, COLLEGE_LINE_3, 12)
    add_centered(document, COLLEGE_LOCATION, 12)
    add_centered(document, DEPARTMENT, 14, True)
    document.add_paragraph()
    add_centered(document, "DECLARATION BY THE CANDIDATES", 16, True)
    add_body(
        document,
        "We, "
        + ", ".join(name for name, _ in TEAM_MEMBERS)
        + ", bearing hall ticket numbers "
        + ", ".join(ticket for _, ticket in TEAM_MEMBERS)
        + f", hereby declare that the project report entitled \"{PROJECT_TITLE}\" is submitted in partial fulfilment of the requirement for the award of the degree of Bachelor of Engineering in Information Technology.",
    )
    add_body(
        document,
        "This report is a record of bonafide work carried out by us. The software implementation, screenshots, testing results, and observations presented in this report have not been submitted to any other university or institute for the award of any other degree or diploma.",
    )
    document.add_paragraph()
    signature_table = document.add_table(rows=4, cols=2)
    signature_table.style = "Table Grid"
    for index, (name, ticket) in enumerate(TEAM_MEMBERS):
        signature_table.cell(index, 0).text = name
        signature_table.cell(index, 1).text = ticket
    signature_table.cell(3, 0).text = "Faculty In-Charge / Reviewer"
    signature_table.cell(3, 1).text = "<Signature and date>"
    document.add_page_break()


def add_acknowledgement_page(document):
    add_heading(document, "ACKNOWLEDGEMENT", 1)
    add_body(
        document,
        "We express our sincere gratitude to the Department of Information Technology, Vasavi College of Engineering, for providing the opportunity to work on this Course Based Project. The project helped us combine frontend interface design, backend service development, and database integration in a single demonstrable application.",
    )
    add_body(
        document,
        f"We thank our faculty coordinators {FACULTY_COORDINATORS[0]} and {FACULTY_COORDINATORS[1]} for their guidance in selecting an appropriate project scope, reviewing progress, and helping us align the implementation with the Full Stack Development course objectives.",
    )
    add_body(
        document,
        "We also acknowledge the official documentation and community resources for HTML5, Canvas API, Node.js, Express, MongoDB, webpack, and dotenv that helped us complete the implementation and documentation in a disciplined manner. Finally, we thank our classmates and reviewers for the feedback that helped refine the usability of the dashboard and the report structure.",
    )
    document.add_page_break()


def add_abstract_page(document):
    add_heading(document, "ABSTRACT", 1)
    add_body(
        document,
        "SYNAPSE Human Connectome Explorer is a course based project that simulates exploration of a human brain connectome through a single page dashboard. The application visualises a sample brain network using the HTML Canvas API, supports region selection, signal propagation, shortest pathway tracing, disease profile switching, selected-region lesioning, and snapshot saving. The frontend is implemented using HTML5, CSS3, and Vanilla JavaScript. The backend is implemented using Node.js and Express, while MongoDB stores the seeded region metadata, network definitions, disease profiles, pathways, and saved simulation snapshots.",
    )
    add_body(
        document,
        "Unlike a purely static mockup, the current version includes a functioning backend API and a working MongoDB data layer. At the same time, the project does not overclaim scientific accuracy. It deliberately uses structured sample data and deterministic graph generation to create a believable and educational connectome dashboard that is practical for a Full Stack Development submission. The project demonstrates UI composition, client-side rendering, backend integration, environment-based configuration, persistence, and reproducible testing in one coherent academic implementation.",
    )
    document.add_page_break()


def add_toc_page(document):
    add_heading(document, "TABLE OF CONTENTS", 1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    run = note.add_run("Update the table of contents in Word or LibreOffice after opening the document to populate page numbers.")
    run.italic = True
    document.add_page_break()


def add_intro_chapter(document):
    add_heading(document, "CHAPTER 1 ABSTRACT AND INTRODUCTION", 1)
    add_heading(document, "1.1 Introduction", 2)
    add_body(
        document,
        "Brain connectivity visualisation is a useful application domain for demonstrating full stack concepts because it naturally combines data organisation, interactive rendering, state management, and persistence. A connectome can be treated as a graph in which regions are nodes and pathways are edges. In SYNAPSE, this idea is adapted into a web application that focuses on usability, interpretability, and presentation readiness rather than real neuroimaging accuracy. The main problem addressed by the project is the lack of an accessible academic dashboard that can explain connectome-style behaviour to a user without requiring specialised medical datasets or tooling [R1][R2][R6].",
    )
    add_body(
        document,
        "The project domain belongs to educational scientific visualisation. The dashboard emulates how a user might inspect brain regions, compare disease-like changes, fire signals through a network, and observe resulting graph-level metrics. The chosen domain makes the application visually rich and technically suitable for a course based submission because it requires frontend design, backend service integration, and document-oriented database storage in the same workflow.",
    )

    add_heading(document, "1.2 Project Objectives and Prioritised Features", 2)
    add_bullets(
        document,
        [
            "Create a polished single page dashboard that looks like a focused application instead of an isolated demo.",
            "Visualise a sample connectome using Canvas-based nodes, edges, animations, and a connectivity matrix.",
            "Support user actions such as region selection, signal firing, shortest pathway tracing, disease profile switching, and lesion simulation.",
            "Integrate MongoDB for storing regions, networks, disease profiles, pathways, and simulation snapshots.",
            "Allow the frontend to remain functional even if the backend is unavailable by using a local fallback dataset.",
            "Improve interpretability through anatomical labels and full-name affected-region displays for disease profiles.",
        ],
    )

    add_heading(document, "1.3 Technology Used", 2)
    add_table(
        document,
        ["Layer", "Technology", "Purpose"],
        [
            ("Frontend structure", "HTML5", "Defines semantic layout for the dashboard, panels, controls, and support sections."),
            ("Styling", "CSS3", "Creates the responsive scientific dashboard look, status chips, cards, and visual states."),
            ("Interaction", "Vanilla JavaScript", "Handles application state, events, graph actions, API calls, and DOM updates."),
            ("Visual rendering", "Canvas API", "Draws the connectome graph, animated signals, labels, and connectivity matrix [R1][R3]."),
            ("Backend", "Node.js + Express", "Serves static assets and exposes REST endpoints for bootstrap and snapshot history [R4][R11]."),
            ("Database", "MongoDB", "Stores seeded connectome-related documents and saved simulation summaries [R5][R6]."),
            ("Tooling", "webpack + npm", "Supports local development, dependency management, and production builds [R8][R9]."),
            ("Configuration", "dotenv", "Loads database URI, database name, and server port through environment variables [R10]."),
        ],
    )

    add_heading(document, "1.4 Software Requirements", 2)
    add_table(
        document,
        ["Component", "Requirement"],
        [
            ("Operating System", "Windows, Linux, or macOS"),
            ("Runtime", "Node.js 20 or later (project verified with Node.js 22)"),
            ("Package Manager", "npm"),
            ("Database", "MongoDB Community Server 8.x or compatible"),
            ("Browser", "Modern browser with Canvas support such as Firefox, Chrome, or Edge"),
            ("Editor / IDE", "VS Code, IntelliJ IDEA, or any text editor"),
        ],
    )

    add_heading(document, "1.5 Hardware Requirements", 2)
    add_table(
        document,
        ["Component", "Minimum", "Recommended"],
        [
            ("Processor", "Dual-core CPU", "Quad-core CPU"),
            ("Memory", "4 GB RAM", "8 GB RAM or above"),
            ("Storage", "2 GB free space", "5 GB free space"),
            ("Display", "1366 x 768", "1920 x 1080"),
            ("Network", "Not mandatory for local demo", "Internet useful for documentation and package installation"),
        ],
    )

    add_heading(document, "1.6 Abstract of the Implemented Solution", 2)
    add_body(
        document,
        "The final implemented solution is a single page application with three main areas: a left panel for selected region and pathway information, a central canvas for connectome visualisation, and a right panel for simulation controls. The application seeds and loads sample region information, creates a weighted graph using deterministic logic, and lets the user observe signal spread, disease-based metric updates, lesion effects, and saved snapshot history. The design intentionally balances academic scope and demonstrable completeness.",
    )


def add_proposed_work_chapter(document):
    add_heading(document, "CHAPTER 2 PROPOSED WORK", 1)
    add_heading(document, "2.1 Problem Statement", 2)
    add_body(
        document,
        "The proposed work was to build a course based project that demonstrates genuine full stack concepts without requiring access to real medical imaging datasets. The challenge was to avoid a superficial prototype while also avoiding unrealistic claims about neuroscience or enterprise-scale infrastructure. The project therefore focuses on a structured simulation architecture: realistic region metadata, deterministic graph generation, visually interpretable actions, an operational backend, and MongoDB persistence for core entities and saved simulation state.",
    )

    add_heading(document, "2.2 Design Overview", 2)
    add_body(
        document,
        "The dashboard design follows a high-density scientific console style. The left panel displays selected region information, activation logs, and path tracing controls. The central canvas acts as the primary visual workspace where nodes, edges, signal ripples, and optional anatomical labels are rendered. The right panel groups disease profiles, metrics, simulation sliders, and action buttons. A matrix footer summarises activity and graph state, while a snapshot section below the dashboard shows saved simulation runs.",
    )
    add_image(document, ASSET_DIR / "synapse_dashboard.png", "Figure 2.1 Main dashboard screenshot of SYNAPSE", width=Inches(6.0))

    add_heading(document, "2.3 Architecture Showing the Flow of Data", 2)
    add_body(
        document,
        "The application follows a browser to API to database flow. On startup, the frontend attempts to load bootstrap data from the Express backend. If the backend is available, MongoDB-backed data is returned. Otherwise, the frontend falls back to a local sample dataset. User actions such as selecting regions, firing signals, changing disease profiles, and tracing pathways are handled in the browser. Saving snapshots uses a POST request that writes summary data into MongoDB.",
    )
    add_image(document, ASSET_DIR / "architecture_diagram.png", "Figure 2.2 Architecture diagram", width=Inches(6.3))

    add_heading(document, "2.4 Activity Diagram", 2)
    add_body(
        document,
        "The main user activity begins with opening the dashboard. The application then loads data, renders the connectome, and waits for interaction. The user can select a region, choose a disease profile, fire a signal, trace a pathway, lesion a region, or save the current snapshot. Each action updates the visual graph, metric values, and history section. Snapshot save is optional and available only when the backend is running.",
    )
    add_image(document, ASSET_DIR / "activity_diagram.png", "Figure 2.3 Activity diagram", width=Inches(6.1))

    add_heading(document, "2.5 Implementation", 2)
    add_body(
        document,
        "The implementation is split into clear modules rather than a single long HTML file. The HTML document defines the structure of the dashboard. The stylesheet contains the dashboard theme, responsive behaviour, and reusable control classes. The JavaScript application layer manages state and event handling. Separate modules define sample connectome data, graph generation logic, and canvas rendering. On the server side, Express routes deliver bootstrap data and receive snapshot requests. MongoDB collections are seeded automatically when the backend starts on an empty database.",
    )

    add_heading(document, "2.6 Module-wise Code and Responsibilities", 2)
    add_table(
        document,
        ["Module / File", "Role in the Project"],
        [
            ("index.html", "Defines topbar, hero, dashboard panels, matrix footer, and snapshot history section."),
            ("css/style.css", "Styles the scientific dashboard interface, cards, controls, labels, panels, and responsive layout."),
            ("js/connectome-data.js", "Stores region blueprints, network definitions, disease profiles, and sample pathways."),
            ("js/connectome-model.js", "Generates nodes and weighted edges, supports path derivation, and locates clicked nodes."),
            ("js/connectome-renderer.js", "Draws graph nodes, edges, matrix, signals, and optional anatomical labels using Canvas."),
            ("js/app.js", "Coordinates UI state, disease switching, label toggling, lesion logic, snapshot loading, and save actions."),
            ("server/server.js", "Starts Express server, serves static files, and defines /api/health, /api/bootstrap, and /api/simulations."),
            ("server/db.js", "Connects to MongoDB, seeds collections, and exposes reusable database helpers."),
            ("server/sample-data.js", "Loads sample data by evaluating the frontend data module for seeding."),
            ("server/seed.js", "Reseeds MongoDB manually through an npm script when needed."),
        ],
    )

    add_heading(document, "2.7 Key Logic Highlighted in Bold", 2)
    p1 = document.add_paragraph()
    p1.add_run("The project uses ").bold = False
    p1.add_run("Seeded Graph Generation").bold = True
    p1.add_run(" to create a stable sample connectome with fixed region metadata but dynamically generated weighted edges based on node distance, same-network grouping, and hub status.")
    p2 = document.add_paragraph()
    p2.add_run("The ").bold = False
    p2.add_run("Signal Ripple Propagation Logic").bold = True
    p2.add_run(" expands concentric ripples from a selected node, activates nearby nodes, and energises strong neighbouring edges.")
    p3 = document.add_paragraph()
    p3.add_run("The ").bold = False
    p3.add_run("Selected-region Lesioning Logic").bold = True
    p3.add_run(" intentionally lesions the currently selected non-hub node, updates metrics, and blocks attempts to lesion protected hub nodes.")
    p4 = document.add_paragraph()
    p4.add_run("The ").bold = False
    p4.add_run("MongoDB Snapshot Persistence Logic").bold = True
    p4.add_run(" writes current disease, selected region, node counts, and metric values into the simulationSnapshots collection for later viewing.")

    add_heading(document, "2.8 Representative Code Snippets", 2)
    add_body(document, "Snippet 1: Express route used to deliver bootstrap data from MongoDB.")
    add_code_block(
        document,
        [
            "app.get(\"/api/bootstrap\", async (_request, response, next) => {",
            "  const [networks, regions, diseaseProfiles, pathways] = await Promise.all([...]);",
            "  response.json({",
            "    source: \"mongodb\",",
            "    NETWORKS: Object.fromEntries(...),",
            "    REGION_BLUEPRINTS: regions,",
            "    DISEASE_PROFILES: Object.fromEntries(...),",
            "    SAMPLE_PATHS: Object.fromEntries(...)",
            "  });",
            "});",
        ],
    )
    add_body(document, "Snippet 2: Canvas label toggle used to reveal anatomical names for main regions.")
    add_code_block(
        document,
        [
            "function toggleLabels() {",
            "  state.labelMode = !state.labelMode;",
            "  elements.toggleLabels.classList.toggle(\"active\", state.labelMode);",
            "  elements.toggleLabels.textContent = state.labelMode",
            "    ? \"Hide Main Region Names\"",
            "    : \"Show Main Region Names\";",
            "}",
        ],
    )

    add_heading(document, "2.9 GitHub Link and Folder Structure", 2)
    add_body(document, f"GitHub repository link: {GITHUB_LINK}")
    add_body(document, "Top-level folder structure used in the project is shown below.")
    add_code_block(
        document,
        [
            "SYNAPSE/",
            "|-- index.html",
            "|-- css/",
            "|   `-- style.css",
            "|-- js/",
            "|   |-- app.js",
            "|   |-- connectome-data.js",
            "|   |-- connectome-model.js",
            "|   `-- connectome-renderer.js",
            "|-- server/",
            "|   |-- server.js",
            "|   |-- db.js",
            "|   |-- seed.js",
            "|   `-- sample-data.js",
            "|-- presentation/",
            "|   `-- SYNAPSE_Presentation.pptx",
            "|-- report_assets/",
            "|-- package.json",
            "|-- package-lock.json",
            "|-- .env.example",
            "|-- webpack.common.js",
            "|-- webpack.config.dev.js",
            "`-- webpack.config.prod.js",
        ],
    )


def add_testing_chapter(document):
    add_heading(document, "CHAPTER 3 TESTING", 1)
    add_body(
        document,
        "Testing for SYNAPSE focused on core user-facing use cases and backend integration points. Each test case was designed to validate whether the application state, visuals, metrics, and persistence behaved as intended under normal and edge-case usage. The project was validated through a combination of manual dashboard interaction, API endpoint checks, build verification, and MongoDB snapshot persistence.",
    )

    headers = ["Test ID", "Use Case", "Test Description", "Expected Result", "Actual Result", "Status"]
    rows = [
        ("TC-01", "Dashboard load", "Open index.html directly from file system.", "Dashboard loads in static sample mode.", "Observed static fallback mode with full interactivity.", "Pass"),
        ("TC-02", "Dashboard load", "Run npm start with MongoDB active and open the app.", "Dashboard shows MongoDB connection status.", "MongoDB Connected indicator displayed.", "Pass"),
        ("TC-03", "Region selection", "Click a visible node in the connectome.", "Selected region card updates with new metadata.", "Name, network, activation, and rank updated.", "Pass"),
        ("TC-04", "Signal propagation", "Press Fire Signal after selecting a node.", "Signal ripple appears and activation metrics change.", "Ripple animation and activation count updated.", "Pass"),
        ("TC-05", "Path tracing", "Select source and target regions and click Trace Pathway.", "Path string, hop count, and strength appear.", "Path details rendered in left panel.", "Pass"),
        ("TC-06", "Disease profile", "Switch from Healthy to Schizophrenia.", "Metrics and affected region tags update.", "Disease metrics and full-name affected regions updated.", "Pass"),
        ("TC-07", "Lesion action", "Select a non-hub node and press Lesion Region.", "Selected node becomes lesioned and metrics drop.", "Node turned red, lesion count increased, metrics decreased.", "Pass"),
        ("TC-08", "Hub protection", "Select a hub node and press Lesion Region.", "Lesion is blocked with guard status.", "Hub Locked message displayed.", "Pass"),
        ("TC-09", "Region labels", "Press Show Main Region Names.", "Main anatomical labels appear on graph.", "Labels for PFC, insula, amygdala, etc. displayed.", "Pass"),
        ("TC-10", "Snapshot save", "Click Save Snapshot with backend active.", "Snapshot saved in MongoDB and history refreshes.", "Snapshot item created and shown in history list.", "Pass"),
        ("TC-11", "API health", "Call GET /api/health.", "API returns ok status and database name.", "API returned JSON with ok=true.", "Pass"),
        ("TC-12", "Build verification", "Run npm run build.", "Production build completes without errors.", "webpack build completed successfully.", "Pass"),
    ]
    add_table(document, headers, rows)


def add_results_chapter(document):
    add_heading(document, "CHAPTER 4 RESULTS", 1)
    add_body(
        document,
        "The implementation produced a working single page dashboard that is demonstrable both as a static preview and as a backend-connected application. Region selection, signal propagation, disease profile switching, selected-region lesioning, label display, and shortest pathway tracing work together in a consistent user flow. With the backend active, the snapshot section also demonstrates full stack persistence by reading and writing to MongoDB.",
    )
    add_body(
        document,
        "The user interface is visually cohesive and presentation-ready. The project does not depend on any paid services or remote infrastructure. It can be demonstrated entirely on a local machine using Node.js, MongoDB Community Server, and a modern browser. This makes the application practical for both evaluation and future extension.",
    )
    add_image(document, ASSET_DIR / "synapse_dashboard.png", "Figure 4.1 Dashboard execution result used for evaluation", width=Inches(6.1))
    add_bullets(
        document,
        [
            "The dashboard loaded successfully in both static mode and backend-connected mode.",
            "MongoDB seeding created the expected collections and sample documents.",
            "Snapshot persistence confirmed that the application is not only visual but also data-driven.",
            "UI improvements such as disease-region tags and anatomical labels made the dashboard easier to interpret during demonstration.",
        ],
    )


def add_knowledge_chapter(document):
    add_heading(document, "CHAPTER 5 ADDITIONAL KNOWLEDGE GAINED", 1)
    add_body(
        document,
        "Implementing SYNAPSE provided knowledge beyond the topics directly covered in a standard Full Stack Development syllabus. The most important additional learning area was Canvas-based data visualisation. Unlike ordinary DOM components, Canvas requires manual control of drawing order, animation timing, coordinate systems, and redrawing logic. This project made those concepts practical by applying them to node-edge graphs, animated ripples, and dynamic text overlays.",
    )
    add_body(
        document,
        "Another major learning outcome was document-oriented data modelling for a simulation-oriented application. Instead of storing only user records or CRUD forms, MongoDB was used to organise structured sample regions, network definitions, disease profiles, pathways, and snapshot summaries. This helped in understanding how a document database can support flexible educational or scientific-style datasets that do not fit neatly into rigid relational tables.",
    )
    add_body(
        document,
        "The project also improved understanding of graceful degradation. The frontend was designed to remain usable even when the backend is unavailable. This required thinking carefully about fallback behaviour, environment-based configuration, and the boundary between persistent data and display-only data. In addition, preparing screenshots, diagrams, and documentation for the project improved our ability to communicate a software system as a complete engineering artefact rather than only as source code.",
    )


def add_conclusion_chapter(document):
    add_heading(document, "CHAPTER 6 CONCLUSION AND FUTURE WORK", 1)
    add_body(
        document,
        "SYNAPSE Human Connectome Explorer successfully meets the objectives of a course based full stack project. It combines a polished scientific-style frontend, Canvas-driven visualisation, application state management, an Express backend, and MongoDB persistence in a demonstrable and coherent system. The final application is truthful about its scope: it is a structured simulation dashboard based on sample data rather than a medical product. Even within that realistic scope, it demonstrates genuine full stack engineering decisions and working integration across multiple layers.",
    )
    add_body(
        document,
        "Future work can improve the project in several directions. The first extension would be importing a larger and more structured connectome dataset while preserving the same interface. The second would be adding authentication and user-specific saved sessions. The third would be expanding the graph logic to support stronger analytical features such as ranking panels, filterable networks, or community detection views. Additional usability improvements such as dark/light themes, exportable reports, or guided tutorial overlays could also make the dashboard more useful for teaching and presentation.",
    )


def add_references_chapter(document):
    add_heading(document, "CHAPTER 7 REFERENCES", 1)
    for key, value in REFERENCE_ITEMS:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.add_run(f"{key}. {value}")


def build_report():
    ensure_assets()
    document = Document()
    set_document_defaults(document)

    add_cover_page(document)
    add_declaration_page(document)
    add_acknowledgement_page(document)
    add_abstract_page(document)
    add_toc_page(document)
    add_intro_chapter(document)
    add_proposed_work_chapter(document)
    add_testing_chapter(document)
    add_results_chapter(document)
    add_knowledge_chapter(document)
    add_conclusion_chapter(document)
    add_references_chapter(document)

    document.core_properties.title = PROJECT_TITLE
    document.core_properties.subject = "Course Based Project Report"
    document.core_properties.author = ", ".join(name for name, _ in TEAM_MEMBERS)
    document.core_properties.comments = f"Generated on {TODAY}"
    document.save(OUTPUT)


if __name__ == "__main__":
    build_report()
